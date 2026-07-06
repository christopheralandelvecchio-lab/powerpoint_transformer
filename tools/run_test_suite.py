"""
run_test_suite.py

Purpose:
    Manifest-driven offline test harness for the PowerPoint Object Extraction
    utility.

Scope:
    This script is included for offline validation by the project owner. It is
    not executed as part of this baseline packaging step.

Design Note:
    The harness exercises the same public package interfaces used by main.py,
    but it compares outputs against manifest-declared expected CSV files.
"""

from __future__ import annotations

import argparse
import csv
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
except ModuleNotFoundError as exc:
    raise ModuleNotFoundError(
        "The test harness Word report requires python-docx. "
        "Install it with: pip install python-docx"
    ) from exc

from src.powerpoint_extractor import (
    extract_connector_relationships,
    extract_pptx_objects,
    write_objects_csv,
    write_relationships_csv,
)


MANIFEST_NAME = "TEST_SUITE_MANIFEST.csv"
TEST_SUITE_DIR = Path(__file__).parent.parent / "test_suite"
ACTUAL_OUTPUT_DIR = Path(__file__).parent.parent / "results"


@dataclass
class TestCase:
    """Manifest row describing one extraction test case."""

    test_case: str
    input_file: Path
    expected_csv: Path
    expected_object_count: int
    expected_relationships_csv: Optional[Path]
    expected_relationship_count: Optional[int]
    purpose: str


@dataclass
class TestResult:
    """Result row written to the offline test reports."""

    test_case: str
    status: str
    actual_count: int
    expected_count: int
    actual_relationship_count: str
    expected_relationship_count: str
    message: str


@dataclass
class TestSummary:
    """Roll-up counts used by the terminal and Word reports."""

    total: int
    passed: int
    failed: int
    errors: int


def normalize_csv_value(value: Any) -> str:
    """
    Normalize a CSV value before comparing expected and actual rows.
    """

    if value is None:
        return ""

    return str(value).strip()


def read_csv_rows(csv_path: Path) -> list[dict[str, str]]:
    """
    Read a CSV file into normalized dictionary rows.
    """

    with csv_path.open("r", newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return [
            {key: normalize_csv_value(value) for key, value in row.items()}
            for row in reader
        ]


def optional_path(TEST_SUITE_DIR: Path, value: str) -> Optional[Path]:
    """
    Convert an optional manifest path cell into a Path or None.
    """

    value = normalize_csv_value(value)

    if not value:
        return None

    return TEST_SUITE_DIR / value


def optional_int(value: str) -> Optional[int]:
    """
    Convert an optional manifest integer cell into an int or None.
    """

    value = normalize_csv_value(value)

    if not value:
        return None

    return int(value)


def load_manifest(TEST_SUITE_DIR: Path) -> list[TestCase]:
    """
    Load test case definitions from TEST_SUITE_MANIFEST.csv.
    """

    manifest_path = TEST_SUITE_DIR / MANIFEST_NAME

    if not manifest_path.exists():
        raise FileNotFoundError(f"Test suite manifest not found: {manifest_path}")

    test_cases: list[TestCase] = []

    with manifest_path.open("r", newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)

        for row in reader:
            test_cases.append(
                TestCase(
                    test_case=row["test_case"],
                    input_file=TEST_SUITE_DIR / row["input_file"],
                    expected_csv=TEST_SUITE_DIR / row["expected_csv"],
                    expected_object_count=int(row["expected_object_count"]),
                    expected_relationships_csv=optional_path(
                        TEST_SUITE_DIR,
                        row.get("expected_relationships_csv", ""),
                    ),
                    expected_relationship_count=optional_int(
                        row.get("expected_relationship_count", ""),
                    ),
                    purpose=row.get("purpose", ""),
                )
            )

    return test_cases


def compare_csv_files(actual_csv: Path, expected_csv: Path) -> tuple[bool, str]:
    """
    Compare two CSV files row-by-row and field-by-field.
    """

    actual_rows = read_csv_rows(actual_csv)
    expected_rows = read_csv_rows(expected_csv)

    if len(actual_rows) != len(expected_rows):
        return False, f"Row count mismatch: actual={len(actual_rows)}, expected={len(expected_rows)}"

    for row_index, (actual, expected) in enumerate(zip(actual_rows, expected_rows), start=1):
        if actual != expected:
            differing_fields = [
                key for key in expected.keys()
                if normalize_csv_value(actual.get(key)) != normalize_csv_value(expected.get(key))
            ]
            return False, f"Row {row_index} mismatch in fields: {', '.join(differing_fields)}"

    return True, "CSV matched expected output."


def run_test_case(test_case: TestCase, ACTUAL_OUTPUT_DIR: Path) -> TestResult:
    """
    Run one manifest-defined test case.
    """

    try:
        objects = extract_pptx_objects(test_case.input_file)
        actual_count = len(objects)

        actual_csv = write_objects_csv(
            objects=objects,
            output_dir=ACTUAL_OUTPUT_DIR,
            pptx_path=test_case.input_file,
        )

        count_matches = actual_count == test_case.expected_object_count
        csv_matches, csv_message = compare_csv_files(actual_csv, test_case.expected_csv)

        relationship_messages: list[str] = []
        actual_relationship_count = ""
        expected_relationship_count = ""

        if test_case.expected_relationships_csv is not None:
            relationships = extract_connector_relationships(objects)
            actual_relationship_count = str(len(relationships))
            expected_relationship_count = str(test_case.expected_relationship_count or 0)

            actual_relationships_csv = write_relationships_csv(
                relationships=relationships,
                output_dir=ACTUAL_OUTPUT_DIR,
                pptx_path=test_case.input_file,
            )

            if len(relationships) != (test_case.expected_relationship_count or 0):
                relationship_messages.append(
                    "Relationship count mismatch: "
                    f"actual={len(relationships)}, expected={test_case.expected_relationship_count}"
                )

            relationships_match, relationship_csv_message = compare_csv_files(
                actual_relationships_csv,
                test_case.expected_relationships_csv,
            )

            if not relationships_match:
                relationship_messages.append(relationship_csv_message)

        if count_matches and csv_matches and not relationship_messages:
            return TestResult(
                test_case=test_case.test_case,
                status="PASS",
                actual_count=actual_count,
                expected_count=test_case.expected_object_count,
                actual_relationship_count=actual_relationship_count,
                expected_relationship_count=expected_relationship_count,
                message="Object CSV and relationship CSV checks passed where applicable.",
            )

        messages = []
        if not count_matches:
            messages.append(
                f"Object count mismatch: actual={actual_count}, expected={test_case.expected_object_count}"
            )
        if not csv_matches:
            messages.append(csv_message)
        messages.extend(relationship_messages)

        return TestResult(
            test_case=test_case.test_case,
            status="FAIL",
            actual_count=actual_count,
            expected_count=test_case.expected_object_count,
            actual_relationship_count=actual_relationship_count,
            expected_relationship_count=expected_relationship_count,
            message=" | ".join(messages),
        )

    except Exception as exc:
        return TestResult(
            test_case=test_case.test_case,
            status="ERROR",
            actual_count=0,
            expected_count=test_case.expected_object_count,
            actual_relationship_count="",
            expected_relationship_count=str(test_case.expected_relationship_count or ""),
            message=f"{type(exc).__name__}: {exc}",
        )


def summarize_results(results: list[TestResult]) -> TestSummary:
    """
    Calculate pass/fail/error totals for the completed test run.
    """

    return TestSummary(
        total=len(results),
        passed=sum(1 for result in results if result.status == "PASS"),
        failed=sum(1 for result in results if result.status == "FAIL"),
        errors=sum(1 for result in results if result.status == "ERROR"),
    )


def write_test_report_csv(results: list[TestResult], ACTUAL_OUTPUT_DIR: Path) -> Path:
    """
    Write offline test results to test_report.csv.

    The CSV report remains useful for automated comparison, filtering, and
    downstream ingestion. The Word report generated separately is intended for
    human review and baseline evidence packages.
    """

    ACTUAL_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    report_path = ACTUAL_OUTPUT_DIR / "test_report.csv"

    with report_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "test_case",
                "status",
                "actual_count",
                "expected_count",
                "actual_relationship_count",
                "expected_relationship_count",
                "message",
            ],
        )
        writer.writeheader()

        for result in results:
            writer.writerow(result.__dict__)

    return report_path


def add_metadata_table(document: Document, TEST_SUITE_DIR: Path, ACTUAL_OUTPUT_DIR: Path) -> None:
    """
    Add run metadata to the Word report.
    """

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    metadata_rows = [
        ("Test suite directory", str(TEST_SUITE_DIR)),
        ("Actual output directory", str(ACTUAL_OUTPUT_DIR)),
        ("Manifest", str(TEST_SUITE_DIR / MANIFEST_NAME)),
    ]

    for label, value in metadata_rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value


def add_summary_table(document: Document, summary: TestSummary) -> None:
    """
    Add pass/fail/error roll-up totals to the Word report.
    """

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Total", "Passed", "Failed", "Errors"]
    values = [summary.total, summary.passed, summary.failed, summary.errors]

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    row = table.add_row().cells
    for index, value in enumerate(values):
        row[index].text = str(value)


def add_results_table(document: Document, results: list[TestResult]) -> None:
    """
    Add detailed per-test results to the Word report.
    """

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.autofit = True

    headers = [
        "Test Case",
        "Status",
        "Objects",
        "Expected Objects",
        "Relationships",
        "Expected Relationships",
        "Message",
    ]

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for result in results:
        row = table.add_row().cells
        row[0].text = result.test_case
        row[1].text = result.status
        row[2].text = str(result.actual_count)
        row[3].text = str(result.expected_count)
        row[4].text = result.actual_relationship_count
        row[5].text = result.expected_relationship_count
        row[6].text = result.message


def write_test_report_docx(
    results: list[TestResult],
    ACTUAL_OUTPUT_DIR: Path,
    TEST_SUITE_DIR: Path,
) -> Path:
    """
    Write offline test results to a Word document.

    The Word report is intended to be readable validation evidence. It includes
    run metadata, summary totals, detailed test results, and a final disposition
    statement that can be reviewed without opening the raw CSV report.
    """

    ACTUAL_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    report_path = ACTUAL_OUTPUT_DIR / "test_report.docx"
    summary = summarize_results(results)

    document = Document()

    section = document.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = document.add_heading("PowerPoint Object Extraction Test Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "This report was generated by the manifest-driven offline test harness. "
        "It summarizes object extraction checks and connector relationship checks "
        "where relationship expectations are declared in the manifest."
    )

    document.add_heading("Run Metadata", level=1)
    add_metadata_table(document, TEST_SUITE_DIR, ACTUAL_OUTPUT_DIR)

    document.add_heading("Execution Summary", level=1)
    add_summary_table(document, summary)

    disposition = "PASS" if summary.failed == 0 and summary.errors == 0 else "FAIL"
    document.add_paragraph(f"Overall disposition: {disposition}")

    document.add_heading("Detailed Results", level=1)
    add_results_table(document, results)

    document.add_heading("Notes", level=1)
    document.add_paragraph(
        "The CSV report remains the machine-readable test artifact. This Word "
        "report is provided as a human-readable companion artifact for review, "
        "baseline evidence, and offline distribution."
    )

    document.save(report_path)
    return report_path


def print_summary(results: list[TestResult], csv_report_path: Path, docx_report_path: Path) -> None:
    """
    Print a concise test harness summary to the terminal.
    """

    summary = summarize_results(results)

    print("\nPowerPoint Object Extraction Test Harness")
    print("=" * 48)

    for result in results:
        print(f"{result.test_case:<8} {result.status:<6} {result.message}")

    print("=" * 48)
    print(f"Total:  {summary.total}")
    print(f"Passed: {summary.passed}")
    print(f"Failed: {summary.failed}")
    print(f"Errors: {summary.errors}")
    print(f"CSV Report:  {csv_report_path}")
    print(f"Word Report: {docx_report_path}")


def parse_args() -> argparse.Namespace:
    """
    Parse test harness command-line arguments.
    """

    parser = argparse.ArgumentParser(description="Run the PowerPoint Object Extraction test suite.")
    parser.add_argument(
        "TEST_SUITE_DIR",
        type=Path,
        help="Directory containing TEST_SUITE_MANIFEST.csv, inputs/, and expected/.",
    )
    parser.add_argument(
        "ACTUAL_OUTPUT_DIR",
        type=Path,
        help="Directory where actual outputs and test_report.csv should be written.",
    )

    return parser.parse_args()


def main() -> None:
    """
    Run all manifest-defined test cases.
    """
    
    test_cases = load_manifest(TEST_SUITE_DIR)
    ACTUAL_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    #args = parse_args()
    #test_cases = load_manifest(args.TEST_SUITE_DIR)
    #args.ACTUAL_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    results = [run_test_case(test_case, ACTUAL_OUTPUT_DIR) for test_case in test_cases]
    csv_report_path = write_test_report_csv(results, ACTUAL_OUTPUT_DIR)
    docx_report_path = write_test_report_docx(
        results=results,
        ACTUAL_OUTPUT_DIR=ACTUAL_OUTPUT_DIR,
        TEST_SUITE_DIR=TEST_SUITE_DIR,
    )
    print_summary(results, csv_report_path, docx_report_path)

    if any(result.status != "PASS" for result in results):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
